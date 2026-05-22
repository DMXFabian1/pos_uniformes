#!/usr/bin/env python3
"""Genera el Tarifario Rápido completo en Word landscape con diseño limpio."""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta ──────────────────────────────────────────────────────────────────
NAVY    = "1E3A6E"
BLUE    = "2E6FBF"
LTBLUE  = "D6E8F7"
RED     = "B52020"
LTRED   = "FCE8E8"
ORANGE  = "C05A00"
LTORANGE = "FEF0E5"
WHITE   = "FFFFFF"
ROWALT  = "EDF4FC"
GRAY    = "F5F5F5"
NOTE    = "666666"

OUTPUT = Path(__file__).parent / "tarifarios"
OUTPUT.mkdir(exist_ok=True)

# ── Helpers XML ─────────────────────────────────────────────────────────────

def _rgb(h: str) -> RGBColor:
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _shade(cell, fill: str) -> None:
    pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def _no_space(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def _cell_margins(cell, top=60, bottom=60, left=100, right=100) -> None:
    pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    pr.append(mar)


def _hide_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        borders.append(b)
    tbl_pr.append(borders)


def _set_col_widths(table, widths_cm: list[float]) -> None:
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j < len(widths_cm):
                cell.width = Cm(widths_cm[j])


def _write_cell(cell, text: str, bold=False, size=10,
                color=None, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = cell.paragraphs[0]
    p.alignment = align
    _no_space(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = _rgb(color)


# ── Layout helpers ───────────────────────────────────────────────────────────

def set_landscape(doc: Document) -> None:
    s = doc.sections[0]
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width  = Cm(29.7)
    s.page_height = Cm(21.0)
    s.left_margin   = Cm(1.3)
    s.right_margin  = Cm(1.3)
    s.top_margin    = Cm(1.2)
    s.bottom_margin = Cm(1.0)


def page_banner(doc: Document, title: str, updated: str = "Abril 2026") -> None:
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _hide_table_borders(t)
    _set_col_widths(t, [20.5, 6.6])

    left = t.rows[0].cells[0]
    right = t.rows[0].cells[1]
    _shade(left, NAVY)
    _shade(right, NAVY)
    _cell_margins(left, top=120, bottom=120, left=200, right=80)
    _cell_margins(right, top=120, bottom=120, left=80, right=200)

    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _no_space(p)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = _rgb(WHITE)

    p2 = right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _no_space(p2)
    run2 = p2.add_run(f"Actualizado: {updated}")
    run2.font.size = Pt(9)
    run2.font.color.rgb = _rgb("9DC3E6")

    spacer = doc.add_paragraph()
    _no_space(spacer)
    spacer.paragraph_format.space_after = Pt(6)


def section_header(doc: Document, text: str,
                   bg: str = BLUE, fg: str = WHITE) -> None:
    t = doc.add_table(rows=1, cols=1)
    _hide_table_borders(t)
    cell = t.rows[0].cells[0]
    _shade(cell, bg)
    _cell_margins(cell, top=80, bottom=80, left=180, right=180)
    _write_cell(cell, text, bold=True, size=11, color=fg)

    spacer = doc.add_paragraph()
    _no_space(spacer)
    spacer.paragraph_format.space_after = Pt(2)


def price_table(doc: Document, headers: list[str],
                rows: list[list[str]],
                col_widths: list[float] | None = None,
                exception: bool = False) -> None:
    hdr_bg  = LTRED   if exception else LTBLUE
    hdr_txt = RED     if exception else NAVY
    alt_bg  = "FDF0F0" if exception else ROWALT

    n = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n)
    t.style = "Table Grid"

    # header row
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]
        _shade(c, hdr_bg)
        _cell_margins(c, top=60, bottom=60, left=120, right=120)
        align = WD_ALIGN_PARAGRAPH.RIGHT if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        _write_cell(c, h, bold=True, size=9, color=hdr_txt, align=align)

    # data rows
    for ri, row in enumerate(rows):
        bg = WHITE if ri % 2 == 0 else alt_bg
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            _shade(c, bg)
            _cell_margins(c, top=55, bottom=55, left=120, right=120)
            is_price = ci > 0
            align = WD_ALIGN_PARAGRAPH.RIGHT if is_price else WD_ALIGN_PARAGRAPH.LEFT
            _write_cell(c, val, bold=is_price, size=10,
                        color=NAVY if is_price else "1A1A2E", align=align)

    if col_widths:
        _set_col_widths(t, col_widths)

    spacer = doc.add_paragraph()
    _no_space(spacer)
    spacer.paragraph_format.space_after = Pt(6)


def note_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph(f"  {text}")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.runs[0]
    run.font.size = Pt(8.5)
    run.font.italic = True
    run.font.color.rgb = _rgb(NOTE)


def columns_layout(doc: Document,
                   slots: list[tuple[str, list[str], list[list[str]]]],
                   n_cols: int = 3,
                   exception: bool = False) -> None:
    """
    Renderiza varias tablas de precios en N columnas usando una tabla contenedora invisible.
    Cada slot: (titulo_seccion, headers, rows)
    """
    hdr_bg  = RED   if exception else BLUE
    hdr_fg  = WHITE
    hdr_light = LTRED if exception else LTBLUE
    alt_bg  = "FDF0F0" if exception else ROWALT

    total_w = 27.1
    col_w = total_w / n_cols

    outer = doc.add_table(rows=1, cols=n_cols)
    _hide_table_borders(outer)
    outer.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i in range(n_cols):
        outer.rows[0].cells[i].width = Cm(col_w)

    for idx, (title, headers, rows) in enumerate(slots):
        if idx >= n_cols:
            # add new row to outer table
            outer.add_row()
            for i in range(n_cols):
                outer.rows[-1].cells[i].width = Cm(col_w)

        row_idx = idx // n_cols
        col_idx = idx % n_cols
        container = outer.rows[row_idx].cells[col_idx]
        _cell_margins(container, top=0, bottom=0, left=60, right=60)

        # section header inside container
        hdr_p = container.add_paragraph()
        _no_space(hdr_p)
        hdr_run = hdr_p.add_run(title)
        hdr_run.bold = True
        hdr_run.font.size = Pt(10)
        hdr_run.font.color.rgb = _rgb(WHITE)
        hdr_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # shade the paragraph's parent (we'll use a small inner table instead)

        # inner table: header bar
        hdr_t = container.add_table(rows=1, cols=1)
        _hide_table_borders(hdr_t)
        hdr_cell = hdr_t.rows[0].cells[0]
        _shade(hdr_cell, hdr_bg)
        _cell_margins(hdr_cell, top=70, bottom=70, left=140, right=140)
        _write_cell(hdr_cell, title, bold=True, size=10, color=WHITE)
        # remove the empty paragraph added earlier
        hdr_p._element.getparent().remove(hdr_p._element)

        # inner price table
        n = len(headers)
        inner = container.add_table(rows=1 + len(rows), cols=n)
        inner.style = "Table Grid"

        for ci, h in enumerate(headers):
            c = inner.rows[0].cells[ci]
            _shade(c, hdr_light)
            _cell_margins(c, top=50, bottom=50, left=100, right=100)
            align = WD_ALIGN_PARAGRAPH.RIGHT if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            _write_cell(c, h, bold=True, size=8.5, color=hdr_bg, align=align)

        for ri, row in enumerate(rows):
            bg = WHITE if ri % 2 == 0 else alt_bg
            for ci, val in enumerate(row):
                c = inner.rows[ri + 1].cells[ci]
                _shade(c, bg)
                _cell_margins(c, top=48, bottom=48, left=100, right=100)
                is_price = ci > 0
                align = WD_ALIGN_PARAGRAPH.RIGHT if is_price else WD_ALIGN_PARAGRAPH.LEFT
                _write_cell(c, val, bold=is_price, size=9,
                            color=NAVY if is_price else "1A1A2E", align=align)

        # small spacer inside container
        sp = container.add_paragraph()
        _no_space(sp)
        sp.paragraph_format.space_after = Pt(4)

    spacer = doc.add_paragraph()
    _no_space(spacer)
    spacer.paragraph_format.space_after = Pt(6)


# ══════════════════════════════════════════════════════════════════════════════
#  PÁGINAS
# ══════════════════════════════════════════════════════════════════════════════

def page_playeras(doc: Document) -> None:
    page_banner(doc, "Playeras Deportivas")
    columns_layout(doc, [
        ("Preescolar", ["Tallas", "Precio"],
         [["2, 4, 6, 8, 10", "$175"], ["12, 14, 16", "$199"], ["CH, MD, GD, EXG", "$219"]]),
        ("Primaria", ["Tallas", "Precio"],
         [["4, 6, 8, 10", "$199"], ["12, 14, 16", "$209"], ["CH, MD, GD, EXG", "$219"]]),
        ("Secundaria", ["Tallas", "Precio"],
         [["12, 14, 16", "$229"], ["CH, MD, GD, EXG", "$239"]]),
        ("Bachillerato", ["Tallas", "Precio"],
         [["12, 14, 16", "$229"], ["CH, MD, GD, EXG", "$239"]]),
    ], n_cols=4)
    note_line(doc, "Si una playera ya tiene precio mayor marcado, se respeta.")


def page_sueteres(doc: Document) -> None:
    page_banner(doc, "Suéteres")
    columns_layout(doc, [
        ("Básico", ["Tallas", "Precio"],
         [["4, 6, 8", "$289"], ["10, 12, 14, 16", "$299"],
          ["30, 32, 34", "$309"], ["36, 38, 40, 42", "$315"], ["44, 46", "$325"]]),
        ("Oficial Preescolar", ["Tallas", "Precio"],
         [["4, 6, 8, 10", "$290"]]),
        ("Oficial Primaria", ["Tallas", "Precio"],
         [["4, 6, 8, 10", "$290"], ["12, 14, 16", "$315"], ["30–42", "$325"]]),
        ("Oficial Secundaria", ["Tallas", "Precio"],
         [["12, 14, 16", "$325"], ["30–44", "$345"]]),
        ("Oficial Bachillerato", ["Tallas", "Precio"],
         [["14, 16", "$325"], ["32, 34", "$335"], ["36, 38", "$345"], ["40, 42, 44", "$355"]]),
    ], n_cols=3)

    columns_layout(doc, [
        ("Excepciones — Preescolar", ["Escuela", "Tallas", "Precio"],
         [["Margarita Paz Paredes", "4, 6, 8, 10", "$315"]]),
        ("Excepciones — Primaria", ["Escuela", "Tallas", "Precio"],
         [["Benito Juárez", "40, 42", "$335"], ["Colegio Motolinea", "40", "$335"],
          ["Francisco Villa", "40", "$335"], ["Ignacio Allende", "40", "$335"],
          ["Miguel Hidalgo", "40, 42", "$335"]]),
        ("Excepciones — Bachillerato", ["Escuela", "Tallas", "Precio"],
         [["UVEG", "14, 16, 32–42", "$355"]]),
    ], n_cols=3, exception=True)

    note_line(doc, "Suéter Oferta Ad hoc Talla Uni se maneja aparte a $199.")
    note_line(doc, "Si una escuela aparece en excepciones, respetar ese precio.")


def page_chamarras(doc: Document) -> None:
    page_banner(doc, "Chamarras")
    columns_layout(doc, [
        ("Deportiva (Con Escudo)", ["Nivel", "Tallas", "Precio"],
         [["Preescolar", "4, 6, 8, 10", "$315"],
          ["Primaria", "4–16, CH, MD, GD", "$335"],
          ["Secundaria", "12–16, CH–EXG", "$385"],
          ["Bachillerato", "12–16, CH–EXG", "$385"]]),
        ("Básica Liso", ["Tallas", "Precio"],
         [["2, 4, 6, 8", "$175"], ["10, 12", "$195"],
          ["14, 16", "$215"], ["28, CH, MD, GD, EXG", "$235"]]),
        ("Básica Punto", ["Tallas", "Precio"],
         [["2, 4, 6, 8", "$275"], ["10, 12, 14, 16", "$295"],
          ["28, CH", "$385"], ["MD", "$395"], ["GD", "$405"], ["EXG", "$415"]]),
    ], n_cols=3)
    note_line(doc, "No heredar precios de Chamarra Básica a Chamarra Deportiva.")
    note_line(doc, "Si una prenda ya tiene precio mayor marcado, se respeta.")


def page_chalecos(doc: Document) -> None:
    page_banner(doc, "Chalecos")
    columns_layout(doc, [
        ("Básico", ["Tallas", "Precio"],
         [["4, 6, 8", "$229"], ["10, 12, 14, 16", "$239"],
          ["28, 30, 32, 34", "$259"], ["36–46", "$269"]]),
        ("Oficial Primaria", ["Tallas", "Precio"],
         [["2, 4, 6, 8", "$240"], ["10, 12", "$245"], ["14, 16", "$250"],
          ["32–42, CH–EXG", "$270"]]),
        ("Oficial Secundaria", ["Tallas", "Precio"],
         [["12, 14, 16", "$245"], ["30, 32, 34", "$265"], ["36–44", "$275"]]),
        ("Oficial Bachillerato", ["Tallas", "Precio"],
         [["14, 16", "$265"], ["32, 34", "$275"], ["36–42", "$285"], ["44", "$295"]]),
    ], n_cols=4)

    columns_layout(doc, [
        ("Excepciones — Primaria (1/2)", ["Escuela", "Tallas", "Precio"],
         [["Huizache", "6, 8", "$240"], ["Huizache", "10, 12", "$245"],
          ["Huizache", "14, 16", "$250"], ["Huizache", "34, 36", "$270"],
          ["Albino García Ramos", "36–42", "$290"]]),
        ("Excepciones — Primaria (2/2)", ["Escuela", "Tallas", "Precio"],
         [["Francisco Villa", "10, 12", "$250"], ["Francisco Villa", "14, 16", "$260"],
          ["Francisco Villa", "34, 36", "$290"], ["Francisco Villa", "38, 40", "$300"],
          ["Justo Sierra", "36", "$275"], ["Justo Sierra", "38", "$280"],
          ["Miguel Campuzano", "34", "$280"], ["Miguel Hidalgo", "36, 38", "$275"],
          ["Miguel Hidalgo", "40, 42", "$285"]]),
    ], n_cols=2, exception=True)

    note_line(doc, "Huizache se normaliza con la escalera de Chaleco Oficial Primaria.")


def page_camisas(doc: Document) -> None:
    page_banner(doc, "Camisas")
    columns_layout(doc, [
        ("Básica Manga Corta Blanca", ["Tallas", "Precio"],
         [["4, 6, 8", "$89"], ["10, 12, 14, 16", "$99"], ["28, 30", "$115"],
          ["32, 34", "$135"], ["36, 38", "$145"], ["40, 42", "$165"]]),
        ("Prowear Manga Corta Blanca", ["Tallas", "Precio"],
         [["4, 6, 8", "$125"], ["10, 12", "$145"], ["14, 16", "$155"],
          ["28–38", "$165"], ["40, 42", "$185"]]),
        ("Manga Larga H Blanca", ["Tallas", "Precio"],
         [["2, 4, 6, 8", "$199"], ["10–18", "$219"], ["32, 34", "$239"],
          ["36, 38", "$249"], ["40, 42, 44, CH", "$259"],
          ["MD", "$269"], ["GD", "$279"], ["EXG", "$289"]]),
        ("Manga Larga M Blanca", ["Tallas", "Precio"],
         [["30, CH", "$269"], ["MD, GD", "$289"], ["EXG", "$309"]]),
        ("Oficial Preescolar", ["Tallas", "Precio"],
         [["Línea completa", "$139"]]),
        ("Oficial Primaria", ["Tallas", "Precio"],
         [["4, 6, 8, 10", "$215"], ["12, 14, 16", "$225"], ["32–38", "$245"]]),
        ("Oficial Secundaria", ["Tallas", "Precio"],
         [["Línea completa", "$265"]]),
        ("Oficial Bachillerato", ["Tallas", "Precio"],
         [["Línea completa", "$265"]]),
    ], n_cols=4)

    columns_layout(doc, [
        ("Excepciones por escuela", ["Escuela", "Línea", "Precio"],
         [["Francisco Villa", "Oficial Primaria", "215 / 225 / 245"],
          ["Bicentenario", "Oficial Secundaria", "$265"],
          ["CBTIS 148", "Oficial Bachillerato", "$265"],
          ["Conalep", "Oficial Escolar", "$265"]]),
    ], n_cols=1, exception=True)

    note_line(doc, "Si aparece Camisa Manga Larga Blanca a $265 sin nivel, conservar precio.")


def page_pants_2pz(doc: Document) -> None:
    page_banner(doc, "Pants 2 Piezas y 3 Piezas")
    columns_layout(doc, [
        ("Preescolar", ["Tallas", "2 pzas", "3 pzas"],
         [["2, 4, 6, 8, 10", "$445", "$545"]]),
        ("Primaria", ["Tallas", "2 pzas", "3 pzas"],
         [["4, 6, 8, 10, 12", "$485", "$585"],
          ["14, 16", "$495", "$595"],
          ["CH, MD, GD", "$585", "$685"]]),
        ("Secundaria", ["Tallas", "2 pzas", "3 pzas"],
         [["10, 12", "$595", "$695"], ["14, 16", "$605", "$705"],
          ["CH, MD, GD, EXG", "$615", "$715"]]),
        ("Bachillerato", ["Tallas", "2 pzas", "3 pzas"],
         [["16", "$615", "$715"], ["CH, MD, GD, EXG", "$615", "$715"]]),
    ], n_cols=4)

    columns_layout(doc, [
        ("Excepciones — Preescolar", ["Escuela", "2 pzas", "3 pzas"],
         [["Sor Juana Inés de la Cruz", "$515", "—"]]),
        ("Excepciones — Primaria", ["Escuela", "2 pzas", "3 pzas"],
         [["Adolfo Lopez Mateo", "575/585/595", "—"],
          ["Colegio Motolinea", "595/630/640/650", "695/730/740/750"]]),
        ("Excepciones — Sec. / Bach.", ["Escuela", "2 pzas", "3 pzas"],
         [["Telesecundaria 454", "$650", "$750"],
          ["Telesecundaria 512", "$650", "$750"],
          ["CECYTE", "$650", "$750"], ["Conalep", "$650", "$750"],
          ["UVEG", "$650", "$750"]]),
    ], n_cols=3, exception=True)


def page_pants_suelto(doc: Document) -> None:
    page_banner(doc, "Pants Suelto Básico")
    columns_layout(doc, [
        ("Liso", ["Tallas", "Precio"],
         [["2, 4, 6, 8", "$165"], ["10, 12, 14, 16", "$180"],
          ["28, CH", "$205"], ["MD", "$215"], ["GD", "$225"], ["EXG", "$235"]]),
        ("Punto", ["Tallas", "Precio"],
         [["2, 4, 6, 8", "$249"], ["10, 12, 14, 16", "$269"],
          ["28, CH", "$309"], ["MD, GD", "$319"], ["EXG", "$339"]]),
    ], n_cols=2)
    note_line(doc, "Liso y Punto son familias distintas — no mezclar escaleras.")


def page_faldas(doc: Document) -> None:
    page_banner(doc, "Faldas")
    columns_layout(doc, [
        ("Gales (Azul, Rojo, Verde)", ["Tallas", "Precio"],
         [["2, 4, 6, 8", "$195"], ["10, 12", "$205"], ["14, 16", "$215"],
          ["28, 30, 32, 34", "$225"], ["36, 38, 40, 42", "$235"], ["44", "$245"]]),
        ("Vino", ["Tallas", "Precio"],
         [["4, 6, 8", "$199"], ["10, 12", "$205"], ["14, 16", "$215"],
          ["28, 30, 32, 34", "$225"], ["36", "$235"]]),
        ("Escocés", ["Tallas", "Precio"],
         [["2, 4, 6, 8", "$205"], ["10, 12, 14, 16", "$215"],
          ["28, 30, 32, 34", "$245"], ["36, 38", "$255"], ["40, 42", "$265"]]),
        ("Gris (Claro y Obscuro)", ["Tallas", "Precio"],
         [["4–16", "$265"], ["28, 30, 32, 34", "$275"],
          ["36, 38", "$285"], ["40, 42", "$295"]]),
        ("Cuatro Tablas", ["Tallas", "Precio"],
         [["2, 4, 6, 8", "$195"], ["10, 12, 14, 16", "$205"],
          ["28, 30, 32, 34", "$225"], ["36, 38, 40, 42", "$235"]]),
        ("Escolar Azul Marino", ["Tallas", "Precio"],
         [["2–16", "$195"], ["18, 20, 28, 30", "$205"],
          ["32, 34", "$215"], ["36, 38", "$225"], ["40, 42", "$235"]]),
        ("Pgallo Café", ["Tallas", "Precio"],
         [["12, 14, 16", "$205"], ["28–34", "$245"], ["36–42", "$255"]]),
        ("Oficial Escolar (genérica)", ["Tallas", "Precio"],
         [["12", "$205"], ["14, 16", "$215"], ["28–34", "$245"], ["36–42", "$255"]]),
    ], n_cols=4)

    columns_layout(doc, [
        ("Excepciones por escuela", ["Escuela", "Tallas", "Precio"],
         [["Colegio Motolinea", "4–8 / 10–12 / 14–16 / 28–34 / 36–38",
           "205/215/225/245/255"],
          ["SABES", "10–16 / 28–34 / 36–38 / 40–42", "270/275/285/295"],
          ["Conalep", "14–16 / 28–34 / 36–42", "280/285/295"],
          ["Margarita Paz Paredes", "4, 6, 8 Preescolar", "$239"]]),
    ], n_cols=1, exception=True)

    note_line(doc, "Las faldas por escuela siempre deben estar por encima de la línea genérica.")


def page_jumper_malla_bata(doc: Document) -> None:
    page_banner(doc, "Jumper · Malla · Bata")
    columns_layout(doc, [
        ("Jumper", ["Tallas", "Precio"],
         [["2, 4, 6, 8", "$249"], ["10, 12, 14, 16", "$269"],
          ["28, 30", "$289"], ["32, 34", "$309"],
          ["36, 38", "$319"], ["40, 42", "$329"]]),
        ("Malla Escolar", ["Tallas", "Precio"],
         [["0-0, 0-2, 3-5, 6-8", "$109"],
          ["9-12", "$129"],
          ["13-18, CH-MD, GD-EXG, Dama", "$139"]]),
        ("Bata", ["Tipo", "Tallas", "Precio"],
         [["Infantil", "Uni", "$65"],
          ["Manga Corta", "12–46", "$395"],
          ["Manga Larga", "12–46", "$439"]]),
    ], n_cols=3)
    note_line(doc, "Si una pieza ya tiene precio mayor marcado, se respeta.")


def page_accesorios(doc: Document) -> None:
    page_banner(doc, "Accesorios · Playera Polo Blanca")
    columns_layout(doc, [
        ("Accesorios", ["Pieza", "Precio"],
         [["Calceta Escolar", "$49"], ["Guante Escolta", "$49"],
          ["Corbata", "$70"], ["Corbatín", "$70"],
          ["Moño", "$70"], ["Boina Escolta", "$80"]]),
        ("Playera Polo Blanca", ["Tallas", "Precio"],
         [["2, 4, 6, 8", "$135"], ["10, 12", "$145"],
          ["14, 16, 18", "$155"], ["28, 30, 32, 34, 36, 38", "$170"],
          ["40, 42", "$180"], ["CH, MD", "$185"], ["GD, EXG", "$195"]]),
    ], n_cols=2)
    note_line(doc, "Si una pieza ya tiene precio mayor marcado, se respeta.")


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════════════

PAGES = [
    ("Playeras",       page_playeras),
    ("Sueteres",       page_sueteres),
    ("Chamarras",      page_chamarras),
    ("Chalecos",       page_chalecos),
    ("Camisas",        page_camisas),
    ("Pants_2pz_3pz",  page_pants_2pz),
    ("Pants_Suelto",   page_pants_suelto),
    ("Faldas",         page_faldas),
    ("Jumper_Malla_Bata", page_jumper_malla_bata),
    ("Accesorios_Polo",   page_accesorios),
]


def generate_all() -> None:
    doc = Document()
    set_landscape(doc)

    # quitar estilo de parrafo por defecto
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(4)

    for i, (name, fn) in enumerate(PAGES):
        if i > 0:
            doc.add_page_break()
        fn(doc)

    out = OUTPUT / "Tarifario_Completo_2026.docx"
    doc.save(str(out))
    print(f"Generado: {out}")

    # también individuales
    for name, fn in PAGES:
        d = Document()
        set_landscape(d)
        st = d.styles["Normal"]
        st.font.name = "Calibri"
        st.font.size = Pt(10)
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after  = Pt(4)
        fn(d)
        p = OUTPUT / f"Tarifario_{name}_2026.docx"
        d.save(str(p))
        print(f"  + {p.name}")


if __name__ == "__main__":
    generate_all()
